# -*- coding: utf-8 -*-
"""生成《企业级私有 Agent 1.0 部署运维文档.docx》到 geesun_agent/docs/。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCX_PATH = r"D:/workspace/geesun_agent/docs/企业级私有Agent1.0部署运维文档.docx"


def _ensure_default_docx():
    """python-docx 内置 templates/default.docx 在隔离 venv 中易损坏（被截断成无效 zip），
    自愈：若默认模板无效，则从同仓 default-docx-template/ 重打包为合法模板。"""
    import zipfile
    import docx
    tpl = os.path.join(os.path.dirname(docx.__file__), "templates", "default.docx")
    src = os.path.join(os.path.dirname(docx.__file__), "templates", "default-docx-template")
    if os.path.isfile(tpl) and zipfile.is_zipfile(tpl):
        return
    if not os.path.isdir(src):
        raise RuntimeError("python-docx 默认模板缺失且无法从 default-docx-template 自愈: " + src)
    with zipfile.ZipFile(tpl, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(src):
            for f in files:
                full = os.path.join(root, f)
                z.write(full, os.path.relpath(full, src))


_ensure_default_docx()
doc = Document()

# ---------- 基础样式 ----------
normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def set_cjk(run):
    run.font.name = "Microsoft YaHei"
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for p in h.runs:
        set_cjk(p)
    return h

def P(text="", bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    set_cjk(run)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_cjk(run)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    # 灰色底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        r = run._element
        r.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        r.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        set_cjk(run)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cjk(run)
            run.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
    run._r.append(fld); run._r.append(instr); run._r.append(fld2); run._r.append(fld3)
    P("（在 Word 中打开后按 F9 更新目录）", italic=True, size=9, color=(120,120,120))

# ================= 封面 =================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("企业级私有 Agent 1.0\n部署与运维文档")
trun.bold = True
trun.font.size = Pt(22)
set_cjk(trun)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("Digital-Factory Brain · geesun_agent 平台")
srun.font.size = Pt(13)
srun.font.color.rgb = RGBColor(0x44,0x44,0x44)
set_cjk(srun)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mrun = meta.add_run("版本 v1.0  ·  适用环境：内网私有化部署  ·  读者：IT 运维 / 接手工程师")
mrun.font.size = Pt(10)
set_cjk(mrun)
doc.add_paragraph()

# ================= 0. 文档说明 =================
H("0. 文档说明", 1)
P("本文档面向内部 IT 运维与新接手工程师，覆盖企业级私有 Agent 1.0（代号 Digital-Factory Brain）从 0 到 1 的部署、组件集成、日常运维与常见故障处理。平台定位为内网安全 + 业务效率的企业级私有 Agent，全部组件私有化部署，不出公网。")
H("0.1 适用范围", 2)
bullet("前端：geesun_agent_web（Next.js 16 + React 19）")
bullet("后端：geesun_agent（FastAPI + LangGraph + DeepAgents）")
bullet("MCP 服务：geesun_mcp_server（FastMCP，含 DLP 解密与沙箱文件通道）")
bullet("沙箱底座：CubeSandbox 0.6.0（MicroVM）+ langchain-cubesandbox 桥接库")
bullet("模型服务：vLLM（Qwen3.6-35B-A3B）")
bullet("可观测：Arize Phoenix + Langfuse（OpenTelemetry 双通道）")
bullet("数据与目录：PostgreSQL（langgraph checkpoint）、LDAP/AD 统一认证")
H("0.2 术语表", 2)
table(
    ["术语", "含义"],
    [
        ["MicroVM", "CubeSandbox 提供的轻量隔离虚拟机，承载代码沙箱执行"],
        ["DLP (ztsm)", "深信服数据防泄漏客户端，文件头 %TSD-Header 标记加密，明文禁止落盘"],
        ["AC-gateway", "内网 SSL 中间人网关，使用伪造 CA 做 HTTPS 审计"],
        ["egress", "出网管控节点（192.168.10.136），仅放行白名单网段，内置 devpi 镜像"],
        ["TTL", "沙箱空闲回收时间（默认 5 分钟）"],
        ["OTLP", "OpenTelemetry 协议，tracing 数据上报通道（gRPC / HTTP）"],
        ["SSE", "Server-Sent Events，后端向浏览器推送流式响应"],
        ["checkpoint", "LangGraph 对话/状态持久化存储（PostgreSQL 或 SQLite）"],
    ],
    widths=[1.6, 5.0],
)

# ================= 1. 系统架构总览 =================
H("1. 系统架构总览", 1)
P("整体为五层拓扑：边缘接入 → 应用层 → 平台/AI 服务 → 沙箱底座 → 基础设施与合规。请求主链路：浏览器经 Nginx 反代访问前端（:3000），前端调用后端 API（:8009），后端编排 LLM、持久化、认证，并通过本地 MCP（:8000）完成 DLP 解密与沙箱文件通道，沙箱底座（CubeAPI :6000）拉起 MicroVM 执行代码。")
H("1.1 组件清单", 2)
table(
    ["组件", "角色", "仓库", "端口", "部署位置"],
    [
        ["geesun_agent_web", "前端 ChatGPT 风格 UI", "geesun_agent_web", "3000", "应用服务器 / Nginx 前"],
        ["geesun_agent", "后端 Agent 编排", "geesun_agent", "8009", "应用服务器"],
        ["geesun_mcp_server", "MCP（DLP解密+沙箱通道）", "geesun_mcp_server", "8000", "与后端同机 localhost"],
        ["vLLM", "大模型推理 Qwen3.6-35B", "独立部署", "8003", "172.16.66.13 裸金属"],
        ["PostgreSQL", "langgraph checkpoint", "独立部署", "5433", "192.168.10.136"],
        ["LDAP/AD", "统一认证", "现有域控", "389", "192.168.1.241"],
        ["Phoenix+Langfuse", "链路追踪", "独立/本地", "OTLP", "可服务端部署"],
        ["CubeSandbox 0.6.0", "沙箱底座（MicroVM）", "CubeSandbox-master", "6000", "172.16.66.13 裸金属"],
        ["langchain-cubesandbox", "Agent→Cube 桥接库", "langchain-cubesandbox", "—", "以 editable 引入后端"],
        ["Harbor", "容器/沙箱镜像仓库", "独立", "8333", "172.16.220.74"],
    ],
    widths=[1.9, 1.9, 1.6, 0.7, 1.5],
)
H("1.2 端口速查", 2)
table(
    ["端口", "服务", "协议", "说明"],
    [
        ["3000", "geesun_agent_web", "HTTP", "Next.js 生产（next start）"],
        ["8009", "geesun_agent", "HTTP", "FastAPI / uvicorn"],
        ["8000", "geesun_mcp_server", "HTTP(streamable)", "仅 127.0.0.1 监听"],
        ["8003", "vLLM", "HTTP/OpenAI", "Qwen3.6-35B-A3B /v1"],
        ["6000", "CubeAPI", "HTTP", "CubeSandbox 控制面"],
        ["5433", "PostgreSQL", "TCP", "agent_mem 库"],
        ["389", "LDAP/AD", "TCP", "目录认证"],
        ["8333", "Harbor", "HTTP", "沙箱镜像仓库"],
        ["13000", "CubeProxy(dev)", "HTTP", "192.168.10.136 开发环境"],
    ],
    widths=[0.8, 2.2, 1.6, 2.4],
)
H("1.3 调用与数据流", 2)
bullet("浏览器 → Nginx(:80/443) → 前端(:3000) → 后端(:8009)")
bullet("后端 → vLLM(:8003) 推理；→ PostgreSQL(:5433) checkpoint；→ LDAP(:389) 认证；→ Phoenix/Langfuse(OTLP) 埋点")
bullet("后端 → MCP(:8000 localhost)：DLP 解密、沙箱文件上传/下载、PDF 比对 stage 编排")
bullet("MCP → DLP 解密 API（%TSD-Header 密文解密）；→ CubeAPI(:6000) 连接 MicroVM")
bullet("CubeAPI → MicroVM 执行代码；Harbor 供给 sandbox-code 镜像")
H("1.4 三仓关系（关键）", 2)
P("后端 geesun_agent 通过 pyproject.toml 的 [tool.uv.sources] 以 editable 本地路径引入 langchain-cubesandbox，因此二者须在构建机同层目录（../langchain-cubesandbox）。geesun_agent_web 与 geesun_mcp_server 为独立仓库，分别部署。skill 目录（含 tech-spec-pdf-diff）位于 geesun_agent/skills/__user_<uid>__/ 下，由 MCP 运行时按全局搜索加载。")

# ================= 2. 前置依赖与环境 =================
H("2. 前置依赖与环境", 1)
H("2.1 硬件与操作系统", 2)
bullet("模型/沙箱底座：Rocky Linux 9 裸金属 @ 172.16.66.13（部署 vLLM + Cubelet 栈）")
bullet("开发/出网节点：OpenCloudOS 9 虚拟机 @ 192.168.10.136（4 vCPU / 7.5Gi，规划扩 16/16Gi）")
bullet("开发机：Windows（D:/workspace 工作区），WSL2 默认 Ubuntu 发行版")
H("2.2 网络与证书", 2)
bullet("egress 192.168.10.136：clean 出网 + devpi PyPI 镜像(:3141) + cube-proxy")
bullet("AC-gateway 172.16.66.13：SSL 中间人，需信任其伪造 CA（cube-root-ca.pem / VeriSign fake）")
bullet("Harbor 172.16.220.74:8333：沙箱镜像源，需提前登录/pull 权限")
bullet("沙箱模板需带 --allow-out-cidr 192.168.10.136/32 放行 devpi 网段")
H("2.3 账号与凭据", 2)
bullet("DLP ztsm 客户端（ztsmdlp.exe 等）须在本机运行，否则加密文件无法解密")
bullet("CubeMaster 凭据（root/Geesun2020.），cube-root-ca.pem 证书")
bullet("LDAP bind 账号（geesunai）+ base DN；JWT_SECRET 随机串")
P("所有凭据通过 .env 注入，严禁入库（.env 已在 .gitignore）。", bold=True)
H("2.4 基础运行时与 SDK 版本要求", 2)
P("本平台由「Python 后端组（3 个组件）+ Node.js 前端组（1 个组件）」构成。部署前须确认目标主机已安装以下基础运行时与包管理器，且版本满足约束，否则依赖安装或前端构建会失败。")
table(
    ["组件", "语言 / 运行时", "版本约束", "包管理器", "说明"],
    [
        ["geesun_agent", "Python", "≥3.13, <3.14（即 3.13.x）", "uv", "后端 FastAPI 服务；依赖 langchain / langgraph / deepagents / pydantic"],
        ["geesun_mcp_server", "Python", "≥3.13（即 3.13.x）", "uv", "MCP 服务（FastMCP）"],
        ["langchain-cubesandbox", "Python", "≥3.13（即 3.13.x）", "uv", "沙箱集成库，以 editable 方式被后端引用（见 pyproject [tool.uv.sources]）"],
        ["geesun_agent_web", "Node.js + bun", "Node 22.x；bun 最新稳定版", "bun", "前端 Next.js 16.2.9 / React 19.2.4（Turbopack）"],
        ["构建机 / CI", "uv + bun", "uv ≥0.5；bun 最新稳定版", "—", "建议统一托管依赖安装与构建，避免版本漂移"],
    ],
    widths=[1.5, 1.6, 2.2, 1.0, 2.0],
)
bullet("三个 Python 组件建议统一使用 Python 3.13.x，并用 uv 管理依赖；geesun_agent 约束最严（<3.14），请勿跨 3.14，也勿降到 3.12 及以下。")
bullet("前端 geesun_agent_web 须 Node.js ≥20.9（生产推荐 22.x LTS），并使用 bun 作为包管理器与运行时；勿用 npm/yarn 替代以免 lockfile 漂移。")
bullet("构建机上 uv 与 bun 须预装：uv 建议 ≥0.5，bun 取最新稳定版；二者均经 egress 192.168.10.136 白名单出网或从内网镜像获取。")
# ================= 3. 逐项组件部署 =================
H("3. 逐项组件部署", 1)

H("3.1 CubeSandbox 0.6.0 底座", 2)
P("裸金属 Rocky Linux 9 @ 172.16.66.13。核心组件：CubeMaster（控制面，root/Geesun2020.）、CubeAPI(:6000)、CubeProxy（openresty 反代）、CubeEgress（MITM，证书 /etc/cube/ca/cube-root-ca.crt）、cube-lifecycle-manager（默认 5 分钟空闲 TTL）。")
P("部署方式（任选其一，来自 CubeSandbox-0.6.0/deploy）：")
bullet("one-click：单主机离线发布包，目标机解包后 install.sh 一键装（推荐首次）")
bullet("kubernetes：容器编排部署")
bullet("pvm / guest-image：裸 MicroVM 镜像或客户机镜像路线")
code("# 单节点配置示例（configs/single-node）\n# cubelet.yaml / cubemaster.yaml / network-agent.yaml\n# NETWORK_CIDR 疑在:\n#   /usr/local/services/cubetoolbox/Cubelet/config/config.toml\n# 镜像源:\n#   sandbox-code:latest  from Harbor 172.16.220.74:8333")
P("已知注意点（部署时核对）：v1.4.9 CPU 过滤 fencepost 曾致 2 核沙箱被拒；cubemaster snapshot_reconciler 在 sandbox_list.go:132 对 GetTraceInfo 返回 nil 做 DeepCopy 会累计 panic；openresty proxy_read_timeout 默认 60s 需调到 300s 并同步 e2b SDK upload gateway timeout，否则大文件上传 504。")

H("3.2 vLLM 模型服务", 2)
code("# 默认模型（OpenAI 兼容 /v1）\nMODEL_NAME=Qwen3.6-35B-A3B\nBASE_URL=http://172.16.66.13:8003/v1")
bullet("vision 仅支持 image_url part；不支持 file type（返回 501 'Unknown part type: file'）——PDF/PPT 须走沙箱侧 pdfplumber/pdftotext 抽文本后再送")
bullet("可动态切换 Kimi / GLM，默认 vLLM Qwen3.6-35B-A3B（MoE 35B）")

H("3.3 PostgreSQL", 2)
code("POSTGRES_HOST=192.168.10.136\nPOSTGRES_PORT=5433\nPOSTGRES_DB=agent_mem\nPOSTGRES_USER=geesun\nPOSTGRES_PASSWORD=<占位符，向 IT 索取>")
P("承载 langgraph checkpoint（postgres / sqlite 双后端支持）。建议每日逻辑备份 + WAL 归档；升级前先 pg_dump agent_mem。")

H("3.4 LDAP/AD 接入", 2)
code("LDAP_SERVER=ldap://192.168.1.241:389\nLDAP_BASE_DN=DC=geesun,DC=li\nLDAP_BIND_USER=geesunai\nLDAP_BIND_PASSWORD=<占位符>\nLDAP_ADMIN_GROUP_DN=CN=geesun-admins,CN=Users,DC=geesun,DC=li\nLDAP_DOMAIN_FORMAT=%s@geesun.li")

H("3.5 geesun_agent 后端", 2)
P("Python >= 3.13，使用 uv 管理。启动命令（start.sh）：")
code("uv run uvicorn src.server:app --host 0.0.0.0 --port 8009 2>&1 | tee server.log")
P("langchain-cubesandbox 以 editable 本地装配：")
code("# geesun_agent/pyproject.toml\n[tool.uv.sources]\nlangchain-cubesandbox = { path = \"../langchain-cubesandbox\", editable = true }")
P(".env 全字段（占位符，真实值向 IT 索取）：")
code("OPENAI_API_KEY=<占位符>\nMODEL_NAME=Qwen3.6-35B-A3B\nBASE_URL=http://172.16.66.13:8003/v1\nEXTRA_MODELS='[]'\n\n# Arize Phoenix 追踪（OTLP gRPC）\nPHOENIX_COLLECTOR_ENDPOINT=<占位符>\n# Langfuse 追踪（OTLP HTTP ingest）\nLANGFUSE_SECRET_KEY=<占位符>\nLANGFUSE_PUBLIC_KEY=<占位符>\nLANGFUSE_BASE_URL=<占位符>\n\nAGENT_WORKSPACE=/mnt/d/workspace/geesun_agent\nUPLOAD_ROOT=/mnt/d/workspace/geesun_agent/data/uploads\nREPORT_ROOT=/mnt/d/workspace/geesun_agent/data/reports\n\n# CubeSandbox\nCUBE_TEMPLATE_ID=<占位符，cubemastercli tpl list 查询>\nCUBE_API_URL=http://172.16.66.13:6000\nCUBE_API_KEY=e2b_0000000000000000000000000000000000000000\nSSL_CERT_FILE=/home/dhp/projects/cube-cert/rootCA.pem\n\n# 数据库 / 认证 / JWT\nPOSTGRES_HOST=192.168.10.136\nPOSTGRES_PORT=5433\nPOSTGRES_DB=agent_mem\nPOSTGRES_USER=geesun\nPOSTGRES_PASSWORD=<占位符>\nLDAP_SERVER=ldap://192.168.1.241:389\nLDAP_BASE_DN=DC=geesun,DC=li\nLDAP_BIND_USER=geesunai\nLDAP_BIND_PASSWORD=<占位符>\nLDAP_ADMIN_GROUP_DN=CN=geesun-admins,CN=Users,DC=geesun,DC=li\nLDAP_DOMAIN_FORMAT=%s@geesun.li\nJWT_SECRET=<随机串，勿用示例>\nJWT_EXPIRE_HOURS=168")
P("tracing 装配依赖：arize-phoenix-otel + openinference-instrumentation-langchain + opentelemetry-exporter-otlp-proto-http，自动对 LangChain/LangGraph 调用埋点。")

H("3.6 geesun_mcp_server", 2)
P("FastMCP 服务，监听 127.0.0.1:8000（streamable-http），与后端同机 localhost 部署。")
code("mcp.run(host=\"127.0.0.1\", port=8000, transport=\"streamable-http\")")
P("关键工具：")
bullet("decrypt_file / decrypt_file_to_base64 / decrypt_and_upload_to_sandbox：DLP 解密（按 %TSD-Header 文件头判定，不靠扩展名）")
bullet("copy_script_to_sandbox / upload_to_sandbox / download_from_sandbox：沙箱文件通道")
bullet("run_pdf_diff_stage1 / run_pdf_diff_stage3：技术协议 PDF 比对确定性流程")
P("环境变量：")
code("DECRYPT_API_URL=<DLP 解密 API 地址>\nE2B_API_URL=<CubeAPI 地址，同 CUBE_API_URL>\nE2B_API_KEY=<占位符>\nSSL_CERT_FILE=<cube-root-ca.pem 路径>\nAGENT_WORKSPACE=<同后端>\nUPLOAD_ROOT=<同后端>\nREPORT_ROOT=<同后端>")

H("3.7 geesun_agent_web 前端", 2)
P("Next.js 16.2.9 + React 19.2.4 + Tailwind + shadcn + @base-ui，Turbopack 构建，使用 bun。")
code("# package.json scripts\n\"dev\": \"next dev\"\n\"build\": \"next build\"\n\"start\": \"next start\"\n\n# .env.local\nNEXT_PUBLIC_API_BASE=http://localhost:8009")
P("生产：next build && next start（默认 :3000），前置 Nginx 反代 :80/:443 到 :3000。注意：WSL 内启动的 dev server 与 Windows 侧 localhost 不通，验证需在 WSL 内或浏览器。")

H("3.8 Arize Phoenix + Langfuse", 2)
P("双通道链路追踪并行采集：")
bullet("Phoenix：OTLP gRPC，由 PHOENIX_COLLECTOR_ENDPOINT 指定 collector 地址")
bullet("Langfuse：OTLP HTTP ingest，由 LANGFUSE_SECRET_KEY/PUBLIC_KEY/BASE_URL 配置")
bullet("openinference-instrumentation-langchain 自动埋点 LangChain/LangGraph，用于调试、评估与回归分析")
P("部署形态：可本地进程或独立服务端；建议独立服务以便长期留存 trace。")

# ================= 4. 集成与联调 =================
H("4. 集成与联调", 1)
H("4.1 端到端冒烟测试", 2)
P("典型链路：上传加密 PDF → 前端 → 后端 → MCP decrypt_and_upload_to_sandbox（解密写沙箱，明文不落盘）→ 沙箱内跑 skill 脚本（extract → diff → generate）→ download_from_sandbox 拉回报告。")
code("# 关键校验点\n# 1) 加密文件上传须走 decrypt_and_upload_to_sandbox（upload_to_sandbox 会拒绝密文）\n# 2) 沙箱内 pdfplumber 缺失时 stage1 自动 pip install（离线需配 devpi 镜像）\n# 3) 扫描件/图片型 PDF（平均页字符 < 50）会明确报错，不静默出空报告")
H("4.2 会话与状态", 2)
bullet("SSE 会话隔离：切换会话存在 race condition 致内容串流的已知 Bug，须使用已修复版本")
bullet("__index__ 索引：更新已有 session 须维护索引，否则 GET /sessions 返回空列表（历史 Bug）")
bullet("沙箱生命周期：业务沙箱由 Agent 管理，GC 不得 kill 业务沙箱（GC-kill 链曾致反复 504）")

# ================= 5. 安全与合规 =================
H("5. 安全与合规", 1)
bullet("DLP 不落盘：解密明文走 decrypt MCP 内存 base64，绝不 open(orig,'wb') 直写；0x80 为完整加密标志、0x00 为半加密态")
bullet("AC-gateway MITM：出网 HTTPS 经伪造 CA，客户端须信任 cube-root-ca.pem")
bullet("egress 管控：仅放行白名单网段；沙箱内 pip 走 devpi 镜像 192.168.10.136:3141")
bullet("JWT：JWT_SECRET 用随机串，有效期 168h；最小权限原则")
bullet(".env 不入库；凭据分级管理，禁止提交明文")

# ================= 6. 运维手册 =================
H("6. 运维手册", 1)
H("6.1 启停", 2)
bullet("后端：./start.sh（uv run uvicorn :8009）或 systemd 托管")
bullet("前端：next build && next start（:3000），Nginx 反代")
bullet("MCP：与后端同机 uv run 启动（:8000）")
bullet("CubeSandbox：systemd 托管 CubeMaster / CubeAPI / CubeProxy / lifecycle-manager")
H("6.2 健康检查", 2)
bullet("端口监听：ss -ltnp | grep -E '8009|8000|6000|8003'")
bullet("后端日志：geesun_agent/server.log（含时间戳，注意 config.py 提前导入曾致无时间戳，已修复）")
bullet("沙箱日志：经 vsock 转宿主 /data/log/CubeShim/；cubemaster.log 在 CubeSandbox 工作区")
H("6.3 监控与告警", 2)
bullet("链路：Phoenix / Langfuse 看 agent 调用链与耗时")
bullet("资源：GPU 显存（vLLM）、MicroVM 数量与空闲（lifecycle TTL）、磁盘（uploads/reports）")
H("6.4 版本升级与回滚", 2)
bullet("代码：git tag 固定版本；多仓变动逐文件 include/exclude 确认后提交")
bullet("镜像：Harbor 镜像版本固定；DB 变更先 pg_dump 再 migration")
bullet("回滚：保留上一镜像标签与 DB 备份，必要时 git revert + 重启")

# ================= 7. 常见 FAQ =================
H("7. 常见故障 FAQ", 1)
faqs = [
    ("上传加密 PDF 报\"检测到 DLP 加密文件\"", "说明文件是密文。必须用 decrypt_and_upload_to_sandbox 解密后传沙箱；upload_to_sandbox 会拒绝密文（密文进沙箱 AI 无法解析）。"),
    ("沙箱跑脚本报 pdfplumber 缺失", "stage1 会自动 pip install pdfplumber；若离线/超时，需在沙箱模板配置 devpi 镜像 192.168.10.136:3141 或预装。"),
    ("模型返回 501 Unknown part type: file", "Qwen API 不支持 file-type content。PDF/PPT 不能直传，须走沙箱侧 pdfplumber/pdftotext 抽文本再送模型。"),
    ("SSE 串会话 / GET /sessions 返回空", "历史 race condition（切换会话内容串流）与 __index__ 索引未维护导致。升级到已修复版本，并确认 session 索引在更新时维护。"),
    ("沙箱调用反复 504", "可能来自 GC-kill 链（沙箱对象 GC 时被 kill）或 openresty proxy_read_timeout 默认 60s 过短。将 proxy_read_timeout 调到 300s 并同步 e2b SDK upload gateway timeout。"),
    ("前端改完代码无变化", "Turbopack 偶发 .next/dev/types/routes.d.ts 缓存失效。删 .next 目录后重启 dev server。"),
    ("解密后仍是明文 PK 未触发 0x80 完整加密", "DLP 策略可能只产出半加密(0x00)。若需完整加密，用 Office 打开后\"另存为\"同路径，Office 保存会触发 DLP 完整加密(0x80)。"),
    ("沙箱内 pip 安装慢/超时", "沙箱内统一走 devpi 镜像 192.168.10.136:3141（模板需 --allow-out-cidr 放行该网段），不要预打包全部依赖。"),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run("Q：" + q); r.bold = True; set_cjk(r)
    pa = doc.add_paragraph()
    ra = pa.add_run("A：" + a); set_cjk(ra)

# ================= 8. 附录 =================
H("8. 附录", 1)
H("8.1 端口速查（汇总）", 2)
table(
    ["端口", "服务", "位置"],
    [
        ["3000", "geesun_agent_web", "应用服务器"],
        ["8009", "geesun_agent", "应用服务器"],
        ["8000", "geesun_mcp_server", "与后端同机"],
        ["8003", "vLLM", "172.16.66.13"],
        ["6000", "CubeAPI", "172.16.66.13"],
        ["5433", "PostgreSQL", "192.168.10.136"],
        ["389", "LDAP/AD", "192.168.1.241"],
        ["8333", "Harbor", "172.16.220.74"],
        ["13000", "CubeProxy(dev)", "192.168.10.136"],
    ],
    widths=[1.0, 2.5, 2.5],
)
H("8.2 仓库与分支索引", 2)
bullet("geesun_agent：后端 + skills/ 目录（tech-spec-pdf-diff 等）")
bullet("geesun_agent_web：前端 Next.js")
bullet("langchain-cubesandbox：Agent→Cube 桥接库（editable 引入后端）")
bullet("geesun_mcp_server：MCP 服务（DLP 解密 + 沙箱通道）")
bullet("CubeSandbox-master：沙箱底座 0.6.0（deploy/、configs/、docs/）")
H("8.3 相关文档", 2)
bullet("各仓 AGENTS.md（开发约束与 lessons-learned）")
bullet("geesun_agent/skills/__user_*/tech-spec-pdf-diff/SKILL.md（PDF 比对 skill 工作流）")
bullet("CubeSandbox-0.6.0/README_zh.md、deploy/one-click/README_zh.md")

# 目录（放最后便于更新；也可放封面后，这里置于文末说明）
doc.add_page_break()
H("目录（更新于文末，Word 中 F9 刷新）", 1)
add_toc()

os.makedirs(os.path.dirname(DOCX_PATH), exist_ok=True)
doc.save(DOCX_PATH)
print("SAVED:", DOCX_PATH, os.path.getsize(DOCX_PATH), "bytes")
